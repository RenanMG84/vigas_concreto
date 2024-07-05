from docx import Document
import os

#Cria um novo documento
res = Document()
 #Add título
res.add_heading("Cálculo de vigas de concreto armado")
res.add_paragraph("")
res.add_paragraph("Parágrafo 1")
res.add_paragraph("")
res.add_paragraph("Parágrafo 2")
res.add_paragraph("")
res.add_paragraph("Parágrafo 3")
res.add_paragraph("")

paragraph = res.add_paragraph("")
paragraph.add_run('Aqui vai o texto 2 !').bold = True

res.add_paragraph("Parágrafo 4")
res.add_paragraph("")

paragraph = res.add_paragraph("")
paragraph.add_run('Aqui vai o texto 3 !').bold = True
res.save("teste.docx")


