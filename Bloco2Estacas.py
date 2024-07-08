import math
from docx import Document
from docx.shared import RGBColor

#DESCRIÇÃO: Cálculo de vigas retangulares em concreto
#Arquivo elaborado com base em BASTOS 2019

#ENTRADA DE DADOS
#Nome do bloco
nome = "B101"
#Força normal característica no pilar
nk = 300 #kN
#Momento fletor característico (no sentido das estacas)
mk = 1000 #kN*cm
#Lado do pilar considerado
ap_pil = 30 #cm
#Diâmetro da estaca
d_est = 30 #cm
#Distância entre eixos de estacas (3xØEst)
e = 80 #cm
#Área do pilar
area_pil = 600 #cm2
#Altura total do bloco
h = 45 #cm
#Maior lado do bloco em planta
l = 150 #cm
#Menor lado do bloco em planta
b = 60 #cm
#Coeficiente de estruturas espaciais
gama_n = 1.2
#Cobrimento
c = 3 #cm
#Diâmetro da armadura principal (5/6.3/8/10/12.5/16/20/25/32)
diam_as_bl = 10 #mm
#Diâmetro da armadura dos estribos verticais e horizontais (5/6.3/8/10/12.5/16)
diam_estr = 10 #mm
#Diâmetro da armadura do pilar (/10/12.5/16/20/25/32)
diam_pil = 10 #mm

#ENTRADA DE DADOS DOS MATERIAIS
fck = 20 #MPa
fyk = 500 #MPa

#TIPO DE SAÍDA DE RELATÓRIO (s = simples / c = completo)
tipo = 'c'






#ALGORITMO----------------------------------------------------------------------------------------------------------------------
#Cria um novo documento word
res = Document()
res.add_heading("Cálculo de blocos de 2 estacas")
res.add_paragraph("")

def qtd_barras(diam, as_calc):
    #Areas de aço por barra (cm2)
    if diam == 5:
        b5 = 0.20
        n_barras = math.ceil(as_calc / b5)
    elif diam == 6.3:
        b63 = 0.315
        n_barras = math.ceil(as_calc / b63)
    elif diam == 8:
        b8 = 0.50
        n_barras = math.ceil(as_calc / b8)
    elif diam == 10:
        b10 = 0.80
        n_barras = math.ceil(as_calc / b10)
    elif diam == 12.5:
        b125 = 1.25
        n_barras = math.ceil(as_calc / b125)
    elif diam == 16:
        b16 = 2
        n_barras = math.ceil(as_calc / b16)
    elif diam == 20:
        b20 = 3.15
        n_barras = math.ceil(as_calc / b20)
    elif diam == 25:
        b25 = 5
        n_barras = math.ceil(as_calc / b25)
    elif diam == 32:
        b32 = 8
        n_barras = math.ceil(as_calc / b32)
    return n_barras

#CONVERSÃO DE VALORES E DEFINIÇÃO DE VARIÁVEIS
fcd = fck / (1.4*10) #kN/cm2
fyd = fyk / (1.15*10) #kN/cm2

#CÁLCULOS--------------------------------------------------------------------------------------
#Força de cálculo majorada
nk = nk * gama_n #kN

#Força normal de cálculo na estaca comprimida
f_est = 1.05*(nk/2) + (mk/e) #kN

#Nova força normal no bloco:
nd = 2 * f_est * 1.4 #kN

#Alturas máximas e mínimas para o bloco, considerando o angulo alpha
d_min = 0.5*(e - ap_pil/2) #cm
d_max = 0.71*(e - ap_pil/2) #cm

#Verificação da altura útil do bloco
d_linha = max(5,(1/5)*(math.sqrt(math.pi)/2)*d_est)
d = h - d_linha #cm

#Ângulo da biela comprimida
alpha = math.atan(d/((e/2)-(ap_pil/4))) #rad

#Verificação das bielas comprimidas junto à seção do pilar e da estaca
#Tensão resistente última na biela (Blevot)
sigma_lim_pil = 1.4*0.95*(fck/(1.4*10)) #kN/cm2
#Tensão solicitante na estaca
ae = (math.pi * math.pow(d_est, 2)) / 4 # cm2 - Área da estaca
sigma_b_est = nd / (2 * ae * math.pow(math.sin(alpha),2)) #kN/cm2
#Tensão solicitante no pilar
sigma_b_pil = nd / (area_pil * math.pow(math.sin(alpha),2)) #kN/cm2

#Cálculo da armadura principal
#As principal - Disposta sobre o topo das estacas
as_calc = (1.15 * nd *(2* e - ap_pil)) / (8 * d * fyd) #cm2
n_barras_as_calc = qtd_barras(diam_as_bl, as_calc)
#Cálculo da armadura superior
as_sup = 0.2 * as_calc #cm2
n_barras_as_sup = qtd_barras(diam_as_bl, as_sup)
#Cálculo da armadura de pele (estribos horizontais) e estribos verticais por face do pilar
as_pele = 0.075 * b * (h/100) #cm2
n_barras_pele = qtd_barras(diam_estr, as_pele)
as_vert = 0.075 * b * (l/100) #cm2
n_barras_vert = qtd_barras(diam_estr, as_vert)
#Espaçamento da armadura de pele
s_pele = min(20, d / 3) #cm
#Espaçamento dos estribos verticais
s_vert = min(15, (0.5 * math.sqrt(math.pi) * d_est) / 2) #cm

#Comprimento de ancoragem das barras no bloco
fctd = 0.15 * math.pow(fck, (2/3)) #MPa
fctd = fctd / 10 #kN/cm2
diam_as_bl = diam_as_bl / 10 #cm
fbd = 2.25 * 1 * 1 * fctd #kN/cm2
lb = (diam_as_bl * fyd) / (4 * fbd) #cm
if lb < 25 * diam_as_bl:
    lb = 25 * diam_as_bl
#Comprimento de ancoragem das barras do pilar
fctd = 0.15 * math.pow(fck, (2/3)) #MPa
fctd = fctd / 10 #kN/cm2
diam_pil = diam_pil / 10 #cm
fbd = 2.25 * 1 * 1 * fctd #kN/cm2
lb_pil = (0.7*(diam_pil * fyd)) / (4 * fbd) #cm
if lb_pil < 25 * diam_pil:
    lb_pil = 25 * diam_pil

#Estimativa do comprimento do bloco
l_min = e - d_est + 2*(0.7 * lb + c + diam_as_bl) #cm
diam_as_calc = diam_as_bl*10 #mm
alpha = math.degrees(alpha)

def res_completo():
    #PREENCHE O DOCUMENTO WORD - COMPLETO-------------------------------------------------------------------
    n1 = res.add_paragraph("")
    n1.add_run('Dados do bloco').bold = True
    res.add_paragraph(f"Bloco: {nome}")
    res.add_paragraph(f"Altura (h): {h} cm")
    res.add_paragraph(f"Maior lado em planta: {l} cm")
    res.add_paragraph(f"Menor lado em planta: {b} cm")
    res.add_paragraph(f"Distância entre estacas: {e} cm")
    res.add_paragraph(f"Diâmetro da estaca: {d_est} cm")
    res.add_paragraph("")
    negrito = res.add_paragraph("")
    negrito.add_run('Resultados ELU').bold = True
    res.add_paragraph(f"Força normal de cálculo nas estacas: {f_est:.2f} kN")
    res.add_paragraph(f"Força normal de cálculo no bloco: {nd:.2f} kN")
    res.add_paragraph("")
    res.add_picture('fig1.png', width=(3000000))
    n2 = res.add_paragraph("")
    n2.add_run('Disposições construtivas').bold = True
    if b < 60 or l < 60:
        p = res.add_paragraph()
        run = p.add_run('Comprimento ou largura menor do que 60 cm!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0)
    else:
        p = res.add_paragraph()
        run = p.add_run('Comprimento e largura maiores do que 60 cm!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0)
    if e < 3*d_est:
        p = res.add_paragraph()
        run = p.add_run('Distância entre estacas menor do que 3Ø!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0)
    else:
        p = res.add_paragraph()
        run = p.add_run('Distância entre estacas maior ou igual 3Ø!')
        font = run.font
        font.color.rgb = RGBColor(0, 100, 0) 
    res.add_paragraph(f"Comprimento mínimo do bloco: e - d_est + 2*(0.7 * lb + c + diam_as_bl) \n \
    {e:.2f} - {d_est:.2f} + 2*(0.7 * {lb:.2f} + {c:.2f} + {diam_as_bl:.2f}) =  {l_min:.0f} cm")
    if l < l_min:
        p = res.add_paragraph()
        run = p.add_run('Comprimento do bloco insuficiente para ancoragem das barras!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('Comprimento do bloco suficiente para ancoragem das barras!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde

    res.add_paragraph(f"Comprimento de ancoragem necessário para barras do pilar: {lb_pil:.0f} cm")
    if d < lb_pil:
        p = res.add_paragraph()
        run = p.add_run('Comprimento do bloco insuficiente para ancoragem das barras do pilar!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('Comprimento do bloco suficiente para ancoragem das barras do pilar!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    res.add_paragraph("")

    n3 = res.add_paragraph("")
    n3.add_run('Dimensões mínimas e ângulo da biela').bold = True
    res.add_paragraph(f"Altura útil mínima do bloco: 0.5*(e - ap_pil/2)  = 0.5*({e} - {ap_pil}/2) = {d_min:.0f} cm")
    res.add_paragraph(f"Altura útil máxima do bloco: 0.71*(e - ap_pil/2) = 0.71*({e} - {ap_pil}/2) = {d_max:.0f} cm")
    res.add_paragraph(f"Altura útil do bloco: {d_max:.0f} cm")
    if d < d_min or d > d_max:
        p = res.add_paragraph()
        run = p.add_run('ALTURA DO BLOCO FORA DOS LIMITES!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('ALTURA DO BLOCO DENTRO DOS LIMITES!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    res.add_paragraph(f"Ângulo da biela comprimida: {alpha:.2f} º")
    if h < d_min or h > d_max:
        res.add_paragraph("ATENÇÃO - ALTURA FORA DOS LIMITES !")
    if alpha < 45 or alpha > 55:
        res.add_paragraph("ATENÇÃO - ÂNGULO FORA DOS LIMITES !")
    res.add_paragraph(f"")
    negrito = res.add_paragraph("")
    negrito.add_run('Verificações na biela').bold = True
    res.add_paragraph(f"Tensão resistente última na biela (Blevot): \n \
    1.4*0.95*(fck/(1.4*10)) = 1.4*0.95*({fck}/(1.4*10)) = {sigma_lim_pil:.2f} kN/cm2")
    res.add_paragraph(f"Tensão solicitante na estaca: nd / (2 * ae * sen(alpha)^2) = \n  {nd:.0f} / (2 * {ae:.0f}* sen({alpha:.2f})^2) = {sigma_b_est:.2F} kN/cm2 \n \
    Área da estaca: ae = (3.1415 * (d_est^2)) / 4  = (3.1415 * {d_est}^2) / 4 = {ae:.2F} cm2")
    if sigma_b_est > sigma_lim_pil:
        p = res.add_paragraph()
        run = p.add_run('TENSÃO NA ESTACA ACIMA DO LIMITE!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('TENSÃO NA ESTACA DENTRO DO LIMITE!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    res.add_paragraph(f"Tensão solicitante no pilar: nd / (area_pil * sen(alpha)^2) \n \
    = {nd:.0f} / ({area_pil:.2f} * sen({alpha:.2f})^2)) = {sigma_b_pil:.2f} kN/cm2")
    if sigma_b_pil > sigma_lim_pil:
        p = res.add_paragraph()
        run = p.add_run('TENSÃO NO PILAR ACIMA DO LIMITE!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('TENSÃO NO PILAR DENTRO DO LIMITE!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    res.add_paragraph("")
    negrito = res.add_paragraph("")
    negrito.add_run('Cálculo das armaduras').bold = True
    res.add_picture('fig2.png', width=(5000000))
    res.add_paragraph(f"* As (arm .principal) - Disposta sobre as estacas: {as_calc:.2f} cm2 - {n_barras_as_calc:.0f} Ø {diam_as_calc:.2f} ")
    res.add_paragraph(f"* Armadura superior (N1): {as_sup:.2f} cm2 - {n_barras_as_sup:.0f} Ø {diam_as_calc:.2f} ")
    res.add_paragraph(f"* Asp - estribos horizontais (por face do bloco): {as_pele:.2f} cm2 - {n_barras_pele:.0f} Ø {diam_estr:.2f} ")
    res.add_paragraph(f"Espaçamento dos estribos horizontais: {s_pele:.0f} cm")
    res.add_paragraph(f"* Asw - estribos verticais (por face do bloco): {as_vert:.2f} cm2 - {n_barras_vert:.0f} Ø {diam_estr:.2f} ")
    res.add_paragraph(f"Espaçamento dos estribos verticais: {s_vert:.0f} cm")
    #Salva o arquivo doc
    res.save("resultados.docx")

def res_simples():
     #PREENCHE O DOCUMENTO WORD - COMPLETO-------------------------------------------------------------------
    negrito = res.add_paragraph("")
    negrito.add_run('Resultados ELU').bold = True
    res.add_picture('fig1.png', width=(3000000))
    n2 = res.add_paragraph("")
    n2.add_run('Disposições construtivas').bold = True
    if b < 60 or l < 60:
        p = res.add_paragraph()
        run = p.add_run('Comprimento ou largura menor do que 60 cm!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0)
    else:
        p = res.add_paragraph()
        run = p.add_run('Comprimento e largura maiores do que 60 cm!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0)
    if e < 3*d_est:
        p = res.add_paragraph()
        run = p.add_run('Distância entre estacas menor do que 3Ø!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0)
    else:
        p = res.add_paragraph()
        run = p.add_run('Distância entre estacas maior ou igual 3Ø!')
        font = run.font
        font.color.rgb = RGBColor(0, 100, 0) 
    if l < l_min:
        p = res.add_paragraph()
        run = p.add_run('Comprimento do bloco insuficiente para ancoragem das barras!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('Comprimento do bloco suficiente para ancoragem das barras!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    if d < lb_pil:
        p = res.add_paragraph()
        run = p.add_run('Comprimento do bloco insuficiente para ancoragem das barras do pilar!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('Comprimento do bloco suficiente para ancoragem das barras do pilar!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    res.add_paragraph("")

    n3 = res.add_paragraph("")
    n3.add_run('Dimensões mínimas e ângulo da biela').bold = True
    if d < d_min or d > d_max:
        p = res.add_paragraph()
        run = p.add_run('ALTURA DO BLOCO FORA DOS LIMITES!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('ALTURA DO BLOCO DENTRO DOS LIMITES!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    if h < d_min or h > d_max:
        res.add_paragraph("ATENÇÃO - ALTURA FORA DOS LIMITES !")
    if alpha < 45 or alpha > 55:
        res.add_paragraph("ATENÇÃO - ÂNGULO FORA DOS LIMITES !")
    res.add_paragraph(f"")
    negrito = res.add_paragraph("")
    negrito.add_run('Verificações na biela').bold = True
    if sigma_b_est > sigma_lim_pil:
        p = res.add_paragraph()
        run = p.add_run('TENSÃO NA ESTACA ACIMA DO LIMITE!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('TENSÃO NA ESTACA DENTRO DO LIMITE!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    if sigma_b_pil > sigma_lim_pil:
        p = res.add_paragraph()
        run = p.add_run('TENSÃO NO PILAR ACIMA DO LIMITE!')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0) #vermelho
    else:
        p = res.add_paragraph()
        run = p.add_run('TENSÃO NO PILAR DENTRO DO LIMITE!')
        font = run.font
        font.color.rgb = RGBColor(0, 128, 0) #verde
    res.add_paragraph("")
    negrito = res.add_paragraph("")
    negrito.add_run('Cálculo das armaduras').bold = True
    res.add_picture('fig2.png', width=(5000000))
    res.add_paragraph(f"* As (arm .principal) - Disposta sobre as estacas: {as_calc:.2f} cm2 - {n_barras_as_calc:.0f} Ø {diam_as_calc:.2f} ")
    res.add_paragraph(f"* Armadura superior (N1): {as_sup:.2f} cm2 - {n_barras_as_sup:.0f} Ø {diam_as_calc:.2f} ")
    res.add_paragraph(f"* Asp - estribos horizontais (por face do bloco): {as_pele:.2f} cm2 - {n_barras_pele:.0f} Ø {diam_estr:.2f} ")
    res.add_paragraph(f"Espaçamento dos estribos horizontais: {s_pele:.0f} cm")
    res.add_paragraph(f"* Asw - estribos verticais (por face do bloco): {as_vert:.2f} cm2 - {n_barras_vert:.0f} Ø {diam_estr:.2f} ")
    res.add_paragraph(f"Espaçamento dos estribos verticais: {s_vert:.0f} cm")
    #Salva o arquivo doc
    res.save("resultados.docx")


if tipo == 'c':
    res_completo()
else:
    res_simples()
    



