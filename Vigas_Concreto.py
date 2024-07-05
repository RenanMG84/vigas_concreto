import math
from docx import Document

#DESCRIÇÃO: Cálculo de vigas retangulares em concreto
# Válido para seções retangulares com fck <=50

#ENTRADA DE DADOS DA VIGA P/ CÁLCULO DAS ARMADURAS
#Nome da viga
nome = "V101"
#Altura total da viga: 
h = 50 #cm
#Largura da viga:
bw = 25 #cm
#Cobrimento
c = 3 #cm
#Diâmetro da armadura escolhida (5/6.3/8/10/12.5/16/20/25/32)
diam = 12.5 #mm
#Momento fletor solicitante de cálculo
msd = 40000 #kN*cm

#ENTRADA DE DADOS DOS MATERIAIS
fck = 20 #MPa
fyk = 500 #MPa

#ENTRADA DE DADOS DA VIGA PARA CÁLCULO DA CORTANTE
#Tipo de modelo de cálculo:
#1 p/ Treliça clássica (ângulo de inclinação das bielas fixo em 45°)
#2 p/ Treliça generalizada (ângulo de inclinação das bielas entre 30º e 45º)
tipo_vsd = 1
#Ângulo de inclinação das bielas - Apenas para modelo 2 (entre 30º e 45º)
teta = 30 #graus
#Cortante solicitante de cálculo
vsd = 300 #kN

#ENTRADA DE DADOS DA VIGA PARA CÁLCULO DA FLECHA
#Momento fletor na seção mais solicitada em serviço (ELS-QP)
mat = 32200 #kN*cm 
#Vão da viga
l_viga = 700 #cm
#Fator que considera deformação por fluência
fl = 2.5
#Carga uniformemente distribuída
q = 8 #kN/m
#Armadura calculada
as_calc = 9.47 #cm2

#ITENS QUE DEVEM SER CALCULADOS (s = sim/ n = nao)
calcular_as_elu = "s"
calcular_cortante = "s"
els_deslocamento = "s"

#ALGORITMO----------------------------------------------------------------------------------------------------------------------
#Cria um novo documento word
res = Document()
res.add_heading("Cálculo de vigas de concreto armado")
res.add_paragraph("")

#Função para escolher o que será calculado
def o_que_calcular(calcular_as_elu, calcular_cortante, els_deslocamento, h, c, diam, msd,bw, fck, fyk, as_calc, mat, q, l_viga, fl, nome, tipo_vsd, teta):
    #CONVERSÃO DE VALORES E DEFINIÇÃO DE VARIÁVEIS
    fcd = fck / (1.4*10) #kN/cm2
    fyd = fyk / (1.15*10) #kN/cm2
    d = h - c - 2 #cm
    res.add_paragraph(f"Viga: {nome}")
    res.add_paragraph(f"Altura (h): {h} cm")
    res.add_paragraph(f"Base (bw): {bw} cm")
    res.add_paragraph("")

    if calcular_as_elu == "s":
        calcular_viga(d, h, msd, bw, fck, fcd, fyd, diam)
    if calcular_cortante == "s":
        global res_vsd
        res.add_heading("Cálculo da força cortante: ") 
        res.add_paragraph("")
        if tipo_vsd == 1:
            calc_vsd_1(bw, d,fck, fcd, vsd)
        else:
            calc_vsd_2(fck, fcd, bw, d, teta, vsd)
    if els_deslocamento == "s":
        els_desl(h, d, bw, fck, msd, as_calc, mat, q, l_viga, fl)

def calcular_viga(d, h, msd, bw, fck, fcd, fyd, diam):
    negrito = res.add_paragraph("")
    negrito.add_run('Resultados para ELU: ').bold = True
    #As mínimo
    as_min = calc_as_min(fck, bw, h)
    as_max = 0.04 * h * bw
    res.add_paragraph(f"Área de aço máxima: {as_max:.2f} cm2 ")
    #Armadura de pele
    if h >= 60:
        as_pele = (0.10 / 100) * (bw * h) #cm2
        res.add_paragraph( f'Armadura de pele: {as_pele:.2f} cm2 por face lateral')
    #Define se é armadura simples ou dupla
    if (math.pow(0.68*d, 2) - 4*0.272*(msd/(bw*fcd))) < 0 :
        calc_as_dupla(h, bw, d, fcd, fyd, msd, diam, as_max)
    else:
        x_1 = (0.68*d + math.sqrt(math.pow(0.68*d, 2) - 4*0.272*(msd/(bw*fcd))))/0.544
        x_2 = (0.68*d - math.sqrt(math.pow(0.68*d, 2) - 4*0.272*(msd/(bw*fcd))))/0.544
        if x_1 > 0 and x_1 <= h:
            x = x_1
        else:
            x = x_2
        if x/d <= 0.45:
            calc_as_simples(msd, fyd,  d, x, as_min, as_max)
        else:   
            calc_as_dupla(h, bw, d, fcd, fyd, msd, diam, as_max)
    
def calc_as_simples(msd, fyd,  d, x, as_min, as_max):
    res.add_paragraph("Cálculo da armadura: Armadura simples")
    res.add_paragraph(f"Profundidade da linha neutra: {x:.2f} cm ")
    as_calc = msd / (fyd*(d - 0.4*x)) #cm2
    res.add_paragraph(f'Área de aço calculada: {as_calc:.2f} cm2')
    #Checa as armaduras mínimas e máximas
    if as_calc < as_min:
        as_calc = as_min
    if as_calc > as_max:
        res.add_paragraph(f"ATENÇÃO: As calculado ({as_calc:.2f} cm2) acima do As Máximo ({as_max:.2f} cm2) permitido!")

    
def calc_as_dupla(h, bw, d, fcd, fyd, msd, diam, as_max):
    res.add_paragraph("Cálculo da armadura: Armadura dupla")
    #Momento limite p/ x = 0,45
    mom_lim = 0.251*bw*math.pow(d,2)*fcd
    #Armadura tracionada - As1
    xlim = 0.45*d
    res.add_paragraph(f'Relação x/d acima de 0,45 - Calculado momento fletor limite M1: {mom_lim:.2f} kN*cm')
    as1 = mom_lim / ((d - 0.4*xlim)*fyd)
    #Armadura tracionada - As2
    m2 = msd - mom_lim
    dlinha = h - d #cm 
    as2 = m2 / ((d - dlinha)*fyd)
    as_tot = as1 + as2
    #qtd barras totais incluindo as barras comprimidas, para o cálculo do As máx
    as_tot_max = as1 + 2* as2
    if as_tot_max > as_max:
        res.add_paragraph(f"ATENÇÃO: As calculado ({as_tot_max:.2f} cm2) acima do As Máximo ({as_max:.2f} cm2) permitido!")
    res.add_paragraph(f'Área de aço p/ armadura tracionada: {as_tot:.2f} cm2 ')
    res.add_paragraph(f'Área de aço p/ armadura comprimida: {as2:.2f} cm2 ')
    res.add_paragraph(f'Quant. barras adotadas p/ armadura tracionada: ')
    qtd_barras(diam, as_tot)
    res.add_paragraph(f'Quant. barras adotadas p/ armadura comprimida: ') 
    qtd_barras(diam, as2)


def calc_vsd_1(bw, d,fck, fcd, vsd):
    #Calculo da cortante pelo modelo de treliça clássica (ângulo de inclinação das bielas fixo em 45°)
    fcd = fcd * 10 #MPa
    res.add_paragraph("Calculo da cortante pelo modelo 1 - Treliça Clássica (ângulo de inclinação das bielas fixo em 45°) ")
    #Verificação de compressão nas bielas
    vrd2 = 0.027*( 1 - (fck/250))*fcd*bw*d #kN
    res.add_paragraph(f"Força cortante solicitante: {vsd:.2f} kN ")
    res.add_paragraph(f"Força resistente a compressão na biela: {vrd2:.2f} kN ")
    if vsd > vrd2:
        res.add_paragraph(f"A força cortante: {vsd:.2f} kN é maior do que a força resistente Vrd2: {vrd2:.2f} kN - NÃO PASSOU !")
    else:
        res.add_paragraph(f'Verificação de compressão nas bielas: OK ')
    #Força cortante referente à armadura mínima
    vsw_min = 0.0137 * bw * d * math.pow(fck, (2/3)) #kN
    if vsd < vsw_min:
        vsd = vsw_min
    res.add_paragraph(f"Força cortante mínima: {vsw_min:.2f} kN")
    #Armadura transversal
    asw_90 = 2.55*(vsd / d) - 0.023 * bw * math.pow(fck, (2/3)) #cm2/m
    res.add_paragraph(f'Área de aço dos estribos: {asw_90:.2f} cm2/m ')
    #Diametros do estribo
    diam_max = bw  #mm
    res.add_paragraph(f'Diâmetro dos estribos: 5 mm à {diam_max:.2f} mm ')
    #Espaçamento máximo dos estribos
    if vsd <= 0.67 * vrd2:
        esp_max = 0.6*d
        if esp_max > 30:
            esp_max = 30
    else:
        esp_max = 0.3*d
        if esp_max > 20:
            esp_max = 20
    res.add_paragraph(f'Espaçamento máximo dos estribos: {esp_max:.2f} cm')

def calc_vsd_2(fck, fcd, bw, d, teta, vsd):
    #Calculo da cortante pelo modelo de treliça clássica (ângulo de inclinação das bielas fixo em 45°)
    global res_vsd
    #converte graus em radianos
    teta = math.radians(teta)
    fcd = fcd * 10 #MPa
    res.add_paragraph("Cálculo pelo Modelo 2 - Treliça generalizada (ângulo de inclinação das bielas entre 30º e 45º)")
    #Força cortante última 
    vrd2 = 0.054 * (1 - fck/250) * fcd * bw * d * math.sin(teta) * math.cos(teta)
    res.add_paragraph(f'Força cortante de cálculo: {vsd:.2f} kN')
    res.add_paragraph(f'Força cortante última: {vrd2:.2f} kN')
    if vsd > vrd2:
        res.add_paragraph("NÃO PASSOU !'- Força cortante maior do que força resistente")
    else:
        res.add_paragraph("Verificação da força cortante: OK, não ocorrerá esmagamento da biela!")
    #Força cortante mínima
    fctd = 0.15 * math.pow(fck, (2/3)) #MPa
    vc0 = 0.6 * (fctd / 10) * bw * d #kN
    vc1 = vc0 * ((vrd2 - vsd) / (vrd2 - vc0)) #kN
    vsd_min = vc1 + 0.0047 * bw * d * math.pow(fck, (2/3)) * (1/ math.tan(teta)) #kN
    res.add_paragraph(f'Força cortante mínima: {vsd_min:.2f} kN ')
    if vsd < vsd_min:
        vsd = vsd_min
    #Armadura transversal
    asw_90 = 2.55*((vsd - vc1) / (d *(1/ math.tan(teta) ))) #cm2/m
    res.add_paragraph(f'Área de aço dos estribos: {asw_90:.2f} cm2/m')
    #Diametros do estribo
    diam_max = bw  #mm
    res.add_paragraph(f'Diâmetro dos estribos: 5 mm à {diam_max:.2f} mm ')
    #Espaçamento máximo dos estribos
    if vsd <= 0.67 * vrd2:
        esp_max = 0.6*d
        if esp_max > 30:
            esp_max = 30
    else:
        esp_max = 0.3*d
        if esp_max > 20:
            esp_max = 20
    res.add_paragraph(f'Espaçamento máximo dos estribos: {esp_max:.2f} cm ')

def els_desl(h, d, bw, fck, msd, as_calc, mat, q, l_viga, fl):
    res.add_heading("Resultados para deslocamento da viga (ELS) ") 
    yt = h /2 #cm
    ic = (bw * math.pow(h, 3)) /12 #cm4
    res.add_paragraph(f"Inércia da seção bruta: {ic:.0f} cm4")
    fctm = (0.3*math.pow(fck, (2/3))) / 10 #kN/cm2
    res.add_paragraph(f"Fctm: {fctm:.2f} kN/cm2")
    mr = (1.5 * fctm * ic) / yt
    res.add_paragraph(f"Momento de fissuração (Mr): {mr:.0f} kN*cm")
    q = q/100 #kN/cm
    if mr < msd:
        res.add_paragraph('Ocorre fissuração!')
        res.add_paragraph('Propriedades da viga no Estádio 2: ')
        alpha_e = 210000/(4760*math.sqrt(fck))
        res.add_paragraph(f'Ae: {alpha_e:.2f} ')
        a1 = bw / 2
        a2 = alpha_e * as_calc
        a3 = - d * alpha_e * as_calc
        res.add_paragraph(f'a1: {a1:.2f} cm ')
        res.add_paragraph(f'a2: {a2:.2f} cm2 ')
        res.add_paragraph(f'a3: {a3:.2f} cm3 ')
        x_1 = (- a2 + math.sqrt(a2 * a2 - 4 * a1 * a3)) / (2 * a1)
        x_2 = (- a2 - math.sqrt(a2 * a2 - 4 * a1 * a3)) / (2 * a1)
        if x_1 > 0 and x_1 <= h:
            x = x_1
        else:
            x = x_2
        res.add_paragraph(f'Profundidade da LN: {x:.2f} cm ')
        iII = ((bw * math.pow(x,3)) / 3) + alpha_e * as_calc * math.pow((x - d), 2)
        res.add_paragraph(f'Inércia da viga fissurada: {iII:.0f} cm4 ')
        im = (math.pow((mr / mat),3))*ic + (1 - math.pow((mr / mat),3))*iII
        res.add_paragraph(f'Inércia média entre Estádios I e II (Branson): {im:.0f} cm4 ')
        i = im
    else:
        res.add_paragraph('Não ocorre fissuração! ')
        i = ic

    #Deslocamento máximo da viga
    res.add_paragraph("Deslocamento máximo da viga: ")
    eci = 5600*math.sqrt(fck) #MPa
    res.add_paragraph(f'Módulo de deformação tangente inicial (Eci): {eci:.0f} MPa ')
    ecs = (0.8 + (0.2*fck)/80) * eci #MPa
    res.add_paragraph(f'Módulo de deformação secante (Ecs): {ecs:.0f} MPa ')
    ecs = ecs /10 #kN/cm2
    desl = ((5 * q * math.pow(l_viga, 4)) / (384*ecs*i)) * fl
    res.add_paragraph(f'Deslocamento máximo da viga: {desl:.2f} cm')
    desl_lim = l_viga / 250
    res.add_paragraph(f'Deslocamento limite da viga: {desl_lim:.2f} cm')
    if desl > desl_lim:
        res.add_paragraph('Não passou!')
    else:
        res.add_paragraph('Passou !')

def calc_as_min(fck, bw, h):
    #AS Mínimo-----------------------------
    if fck == 20 or fck == 25 or fck == 30:
        p = 0.15 / 100
    elif fck == 35:
        p = 0.164 /100
    elif fck == 40:
        p = 0.179/100
    elif fck == 45:
        p = 0.194/100
    elif fck == 50:
        p = 0.208/100
    as_min = p*bw*h
    res.add_paragraph(f'Armadura mínima: {as_min:.2f} cm2 ')
    return as_min

def qtd_barras(diam, as_calc):
    #Areas de aço por barra (cm2)
    if diam == 5:
        b5 = 0.20
        n_barras = math.ceil(as_calc / b5)
    elif diam == 6.3:
        b63 = 0.315
        n_barras = math.ceil(as_calc / b63)
    elif diam == 8:
        b8 = 0.50
        n_barras = math.ceil(as_calc / b8)
    elif diam == 10:
        b10 = 0.80
        n_barras = math.ceil(as_calc / b10)
    elif diam == 12.5:
        b125 = 1.25
        n_barras = math.ceil(as_calc / b125)
    elif diam == 16:
        b16 = 2
        n_barras = math.ceil(as_calc / b16)
    elif diam == 20:
        b20 = 3.15
        n_barras = math.ceil(as_calc / b20)
    elif diam == 25:
        b25 = 5
        n_barras = math.ceil(as_calc / b25)
    elif diam == 32:
        b32 = 8
        n_barras = math.ceil(as_calc / b32)
    res.add_paragraph( f'Qtd de barras de {diam:.2f} mm: {n_barras:.0f} ')

def ancoragem():
    pass

def torcao():
    pass

o_que_calcular(calcular_as_elu, calcular_cortante, els_deslocamento, h, c, diam, msd,bw, fck, fyk, as_calc, mat, q, l_viga, fl, nome, tipo_vsd, teta)

#Salva o arquivo doc
res.save("resultados.docx")






